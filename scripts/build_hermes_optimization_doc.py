from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\Yu\Desktop\tianjun-optimize\tianjun-design document\Hermes使用设计评估与优化说明-最终版.docx")

NAVY = "163A5F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5A6470"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "2E6B4D"
GOLD = "7A5A00"
RED = "9B1C1C"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, *, color: str = "D7DEE7", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is not None:
        return
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_font(run, *, name: str = "Calibri", east_asia: str = "Microsoft YaHei", size: float = 11,
             bold: bool | None = None, italic: bool | None = None, color: str = BLACK) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_style_font(style, *, size: float, color: str = BLACK, bold: bool | None = None) -> None:
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold


def set_style_sheet(doc: Document) -> None:
    normal = doc.styles["Normal"]
    apply_style_font(normal, size=10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        apply_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color=MUTED)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)
    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    left = header.add_run("TIANJUN ENGINE  /  HERMES OPTIMIZATION REPORT")
    set_font(left, size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    add_page_field(footer)


def add_para(doc: Document, text: str = "", *, size: float = 10.5, color: str = BLACK,
             bold: bool = False, italic: bool = False, before: float = 0, after: float = 6,
             align=WD_ALIGN_PARAGRAPH.LEFT, keep_with_next: bool = False) -> object:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_rich_para(doc: Document, parts: Iterable[tuple[str, dict]], *, after: float = 6) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.1
    for text, kwargs in parts:
        run = paragraph.add_run(text)
        set_font(run, **kwargs)
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_callout(doc: Document, label: str, text: str, *, fill: str = CALLOUT, accent: str = NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(f"{label}  ")
    set_font(run, size=10.2, bold=True, color=accent)
    run = paragraph.add_run(text)
    set_font(run, size=10.2, color=BLACK)
    add_para(doc, "", after=3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int],
              *, header_fill: str = LIGHT_BLUE, font_size: float = 9.2) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_font(run, size=font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            if row_index % 2 == 0:
                set_cell_shading(cell, "FBFCFD")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.03
            run = paragraph.add_run(value)
            set_font(run, size=font_size, color=BLACK)
        set_table_geometry(table, widths_dxa, indent_dxa=120)
    add_para(doc, "", after=3)
    return table


def add_numbering(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract or [0]) + 1
    num_id = max(existing_num or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc: Document, text: str, *, num_id: int, size: float = 10.3) -> object:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(num)
    p_pr.append(num_pr)
    run = paragraph.add_run(text)
    set_font(run, size=size)
    return paragraph


def add_cover(doc: Document) -> None:
    add_para(doc, "TIANJUN ENGINE", size=10, color=BLUE, bold=True, after=26)
    add_para(doc, "Hermes 使用设计评估与优化说明", size=27, color=NAVY, bold=True, after=8)
    add_para(doc, "面向算力网络智能调度的意图理解、反馈闭环与安全工具链优化", size=14, color=DARK_BLUE, after=22)
    add_callout(
        doc,
        "交付说明",
        "本报告基于 hermes-2.txt 的优化建议与当前天钧引擎源码核验结果编写。报告区分既有能力、本次实际落地优化和后续研究方向，并记录可复现测试结果。",
        fill=LIGHT_BLUE,
    )
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["评估对象", "天钧引擎中的 Hermes 受控智能体使用链路"],
            ["输入材料", r"C:\Users\Yu\Desktop\hermes-2.txt 与当前工作区源码"],
            ["优化范围", "需求解析、地域识别、优先级建模、置信度、反馈调权、多轮追问"],
            ["验证方式", "源码核验、Python 编译检查、pytest 回归测试"],
            ["文档日期", "2026 年 6 月 2 日"],
        ],
        [1700, 7660],
        header_fill=LIGHT_GRAY,
        font_size=9.6,
    )
    add_para(doc, "", after=12)
    add_para(doc, "结论预览", size=12, color=BLUE, bold=True, after=6)
    add_para(
        doc,
        "Hermes 原有架构已经具备“LLM 辅助理解、确定性控制面决策、人工确认后提交”的关键安全边界。本次优化没有重复建设已有的 LLM 兜底，而是补齐了确定性解析、连续意图参数、反馈权重闭环和主动澄清能力，使自然语言意图能更稳定地进入真实调度评分。",
        size=11,
        color=NAVY,
        after=12,
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build_document() -> None:
    doc = Document()
    set_style_sheet(doc)
    configure_page(doc)
    bullets = add_numbering(doc, ordered=False)
    numbers = add_numbering(doc, ordered=True)
    research_numbers = add_numbering(doc, ordered=True)

    add_cover(doc)

    add_heading(doc, "1. 评估结论", 1)
    add_para(
        doc,
        "hermes-2.txt 对当前系统的判断总体方向正确，尤其准确识别了反馈解析与真实调度指标之间的命名脱节，以及枚举型 priority 无法表达复合目标的问题。但源码核验也表明，部分建议描述的是系统已经具备的能力：ChatRuntime 已接入 OpenAI-compatible LLM，并以受控槽位跟踪器方式调用 DeepSeek；LLM 只提出结构化更新，控制面仍负责校验、策略生成、仿真和提交保护。",
    )
    add_callout(
        doc,
        "本次策略",
        "保留现有 LLM-first 但非 LLM-authoritative 的架构，强化离线可运行的确定性底座。即使未配置模型、模型调用失败或用户输入很短，Hermes 仍可给出可审计的需求结构、追问和调度权重。",
    )
    add_table(
        doc,
        ["方向", "文本判断", "源码核验", "本次处理"],
        [
            ["混合解析层", "建议正则优先、LLM 低置信度兜底", "LLM 槽位跟踪已存在，但确定性长尾覆盖不足", "扩展隐式业务识别，并保留现有 LLM 状态跟踪"],
            ["地域识别", "别名较少，武汉与服务区域存在脱节", "判断成立", "扩充城市与缩写、加入编辑距离纠错、纳入 wuhan 服务区域"],
            ["优先级向量", "枚举无法表达“低时延但不超预算”", "判断成立", "新增连续 priority_vector，映射到真实评分维度"],
            ["置信度", "字段数量加分无法反映调度影响", "判断成立", "改为 workload-aware 的 per-slot 加权置信度"],
            ["反馈解析", "反馈 delta 键与 METRIC_KEYS 不一致", "判断成立，属于真实闭环脱节", "统一到九项真实指标，并增加程度词幅度"],
            ["主动追问", "缺失字段按固定顺序追问", "判断成立", "按影响启发式排序，并先询问不可用地域是否可放宽"],
        ],
        [1300, 2300, 2400, 3360],
        font_size=8.7,
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "2. Hermes 原有使用设计", 1)
    add_para(
        doc,
        "Hermes 是天钧引擎面向用户和智能体宿主的受控交互层。它不是一个可以直接修改集群状态的自由代理，而是围绕唯一可信控制面组织对话、工具调用和人工确认。Dashboard、进程内 ChatRuntime 与 MCP adapter 共用 TianjunToolService 契约，减少入口之间的行为漂移。",
    )
    add_table(
        doc,
        ["层次", "核心模块", "职责", "安全边界"],
        [
            ["对话层", "ChatRuntime", "识别用户意图、管理会话、调用需求澄清与策略工具", "LLM 只能辅助解析与组织回复"],
            ["工具层", "TianjunToolService / MCP", "提供查询、澄清、草拟、仿真、解释、优化、提交和调度工具", "状态变化必须走显式工具"],
            ["控制面", "CentralControlPlane", "维护节点、任务、会话、策略、lease、执行反馈与报告", "库存事实只来自注册节点或外部上报"],
            ["策略层", "ComputeNetworkPolicyGenerator", "把需求转为任务画像、生成候选策略、解释风险", "候选策略与仿真结果可审计"],
            ["调度层", "ClosedLoopAdaptiveScheduler", "九项指标归一化评分、硬约束过滤、确定性选点", "最终选点不交给 LLM"],
            ["执行层", "Node Agent / Simulation / CloudSimPlus", "执行任务并回传进度、耗时、成本和错误", "正式下发前需要显式确认"],
        ],
        [1050, 1900, 3200, 3210],
        font_size=8.8,
    )
    add_heading(doc, "2.1 标准使用流程", 2)
    for item in (
        "用户通过 Dashboard、HTTP 或 MCP Host 描述自然语言需求。",
        "Hermes 调用 start_requirement_dialogue 或 continue_requirement_dialogue，将文本转为结构化槽位，并根据缺失项追问。",
        "需求可用后，Hermes 调用 draft_compute_network_policy 与 simulate_policy，输出推荐节点、稳定时延、预算、安全和风险解释。",
        "用户通过确认参数或 Dashboard 的正式下发动作授权提交。聊天中的普通“确认”不会绕过保护。",
        "任务创建 lease，由仿真后端或真实节点执行，并回传进度与结果。",
        "控制面根据执行结果更新节点健康度、性能因子和策略权重，Dashboard 与 /report 展示闭环效果。",
    ):
        add_list_item(doc, item, num_id=numbers)

    add_heading(doc, "3. 本次实际落地的优化", 1)
    add_heading(doc, "3.1 确定性隐式意图识别", 2)
    add_para(
        doc,
        "原有解析器可以识别“推理”“在线服务”“分析”等显式表达，但对业务化说法覆盖有限。本次扩充了无需 LLM 即可命中的表达，包括在线问答、问答服务、对话服务、embedding、向量化、向量嵌入、批量打标签、批量标注、数据清洗和 ETL。这样可以降低模型不可用时的错误回落率，并让规范输入和长尾输入共享同一 UserRequirement 结构。",
    )
    add_table(
        doc,
        ["用户表达", "优化前风险", "优化后识别"],
        [
            ["搞个在线问答服务", "可能回落到 batch", "inference"],
            ["帮我跑个 embedding", "可能回落到 batch", "inference"],
            ["批量打标签", "可能被当作普通 batch", "analytics"],
        ],
        [2900, 3300, 3160],
        font_size=9.2,
    )

    add_heading(doc, "3.2 地域字典扩展与模糊纠错", 2)
    add_para(
        doc,
        "地域解析新增华东、西部、华南多个城市别名，并支持 cd、cq 等短缩写。短缩写使用单词边界匹配，避免在普通英文单词中误触发。对带上下文的中文地域表达，解析器加入编辑距离不超过 1 的模糊纠错，例如“部署在成嘟”可恢复为 west。武汉被纳入 SERVICE_REGION_CODES，使其不再被误当作数据源物理端点。",
    )
    add_table(
        doc,
        ["类别", "新增示例", "处理方式"],
        [
            ["华东", "南京、苏州、无锡、宁波、合肥、济南、青岛、天津", "统一映射 east"],
            ["西部", "西安、昆明、贵阳、兰州、乌鲁木齐、cd、cq", "统一映射 west"],
            ["华南", "厦门、福州、南宁、海口", "统一映射 south"],
            ["华中", "武汉 / wuhan", "独立 service region code"],
            ["纠错", "成嘟", "在地域上下文中用编辑距离恢复成都"],
        ],
        [1200, 4100, 4060],
        font_size=9.1,
    )

    add_heading(doc, "3.3 从枚举 priority 到连续 priority_vector", 2)
    add_para(
        doc,
        "为兼容现有接口，priority 枚举仍保留为主目标；新增 priority_vector 表达复合目标。解析器会同时提取 latency、cost、quality、security、balance、fragmentation、locality 和 network 八类意图，并归一化为连续向量。任务生成时，该向量被映射到调度器实际使用的九项指标增量。",
    )
    add_table(
        doc,
        ["意图维度", "映射到调度指标", "作用"],
        [
            ["latency", "performance、completion、network", "偏向响应更快、网络更稳的节点"],
            ["cost", "cost", "提高成本得分影响"],
            ["quality", "reliability、completion、performance", "偏向 SLA 与完成质量"],
            ["security", "security、reliability、locality", "加强隔离、可靠性与驻留"],
            ["balance", "balance", "降低热点节点压力"],
            ["fragmentation", "fragmentation", "减少稀缺资源碎片"],
            ["locality", "locality", "提高就近与数据驻留偏好"],
            ["network", "network、performance", "提高网络画像影响"],
        ],
        [1500, 3400, 4460],
        font_size=9.0,
    )
    add_callout(
        doc,
        "创新价值",
        "“低时延但不能超预算”不再被压扁成一个枚举。自然语言中的多目标权衡会进入 task.intent_weights，再参与 _derive_task_weights()，形成从意图理解到调度评分的完整可观测链路。",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "3.4 Workload-aware 的槽位置信度", 2)
    add_para(
        doc,
        "原有 confidence 采用固定加分与缺失扣分规则。本次增加 slot_confidence，并按任务类型调整槽位权重：推理与流式任务提高时延槽位的重要性，训练任务提高资源规格的重要性。明确填写、允许使用默认值和关键字段缺失会得到不同分值，使 confidence 更接近“当前需求对调度结果有多确定”。",
    )
    add_table(
        doc,
        ["槽位", "典型显式值", "默认或可推断值", "关键缺失值"],
        [
            ["workload_type", "1.00", "0.76", "0.00"],
            ["region_preference", "1.00", "明确不限时 0.72", "0.00"],
            ["resources", "1.00", "0.62", "按工作负载默认画像"],
            ["latency_target_ms", "1.00", "非关键时 0.58", "低时延意图下 0.00"],
            ["budget_limit", "1.00", "非关键时 0.58", "成本意图下 0.00"],
        ],
        [1900, 1900, 2760, 2800],
        font_size=9.1,
    )

    add_heading(doc, "3.5 反馈语义与真实调度权重对齐", 2)
    add_para(
        doc,
        "hermes-2.txt 指出的 bug 已修复：反馈解析不再生成调度器无法消费的 latency_weight、quality_weight 等孤立键，而是直接输出 METRIC_KEYS 中的真实维度。为兼容旧调用方，旧键仍会在入口处归一化到新键。用户反馈可以影响均衡、碎片、本地性和网络，不再只覆盖时延、成本、质量和安全。",
    )
    add_table(
        doc,
        ["反馈表达", "结构化目标", "写入真实指标"],
        [
            ["响应太慢", "latency", "performance、completion、network"],
            ["预算太高", "cost", "cost"],
            ["SLA 不稳定并且丢包", "qos", "reliability、completion、network"],
            ["节点负载太高", "balance", "balance"],
            ["GPU 排队且碎片严重", "fragmentation", "fragmentation"],
            ["数据不在本地", "locality", "locality"],
            ["网络抖动明显", "network", "network"],
        ],
        [2500, 1600, 5260],
        font_size=9.0,
    )
    add_heading(doc, "3.6 幅度感知反馈", 2)
    add_para(
        doc,
        "反馈调整不再固定为 0.22。解析器会根据程度词选择幅度：稍微或略微为 0.05，有点或偏为 0.12，默认反馈为 0.18，太或明显为 0.22，极度、严重、完全无法接受为 0.35。该设计为后续实验提供了连续、可解释的用户偏好更新入口。",
    )
    add_table(
        doc,
        ["程度", "示例", "调整幅度"],
        [
            ["轻微", "响应稍微慢了一点", "0.05"],
            ["较轻", "成本有点高", "0.12"],
            ["一般", "网络不稳定", "0.18"],
            ["明显", "节点负载太高", "0.22"],
            ["严重", "响应完全无法接受", "0.35"],
        ],
        [1600, 5560, 2200],
        font_size=9.2,
    )

    add_heading(doc, "3.7 调度影响导向的多轮追问", 2)
    add_para(
        doc,
        "多轮澄清从固定 missing_fields 顺序升级为可解释的影响启发式排序。地域、工作负载、时延、预算和安全槽位有不同基础影响；推理与流式任务会提高时延问题优先级，训练任务会提高地域问题优先级，priority_vector 也会继续修正排序。当用户指定的地域没有在线节点时，Hermes 会优先询问是否放宽地域或提供备选地域。",
    )
    add_callout(
        doc,
        "研究边界",
        "当前实现属于低成本、可解释的主动澄清启发式。真正的信息增益最大化仍可作为后续论文方向：对候选槽位做反事实填充，比较 select_node() 排名变化和可行候选集合变化，再决定下一问。",
        fill="FFF8E8",
        accent=GOLD,
    )

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    add_heading(doc, "4. 完整 Hermes 创新点与功能", 1)
    add_para(
        doc,
        "在本次优化后，Hermes 的价值不局限于自然语言入口。它把意图理解、约束澄清、确定性选点、模型增强、仿真验证、人工授权和执行反馈组织成一条可审计闭环。",
    )
    innovations = [
        ("LLM-first 但非 LLM-authoritative", "DeepSeek 可做对话状态追踪和槽位更新建议；库存事实、状态迁移、选点与提交仍由确定性控制面负责。"),
        ("确定性与 LLM 双路径意图理解", "高频表达由本地规则快速处理，复杂上下文由 LLM 辅助；模型不可用时仍能运行。"),
        ("连续多目标意图向量", "把时延、成本、质量、安全、均衡、碎片、本地性和网络偏好映射到真实调度权重。"),
        ("算力与网络联合决策", "选点同时考虑 CPU/GPU、成本、可靠性、负载、碎片、局部性、网络和安全。"),
        ("拓扑感知模型增强", "LSTM 与 GraphSAGE 可融合进入网络稳定性评分；缺少模型或分布外输入时透明降级。"),
        ("人工确认工具边界", "Hermes 可以草拟、仿真、解释和优化，但正式下发必须显式授权。"),
        ("反馈驱动双闭环", "执行记录驱动 PolicyOptimizer 自动调权，用户自然语言反馈驱动偏好权重更新。"),
        ("统一工具面", "Dashboard、进程内 Hermes、HTTP 与 MCP adapter 共用 TianjunToolService 契约。"),
        ("可观测与可追溯", "策略对象保留需求、权重、网络快照、拓扑证据、预期效果、风险与解释。"),
    ]
    add_table(
        doc,
        ["创新点", "功能说明"],
        [[name, description] for name, description in innovations],
        [2600, 6760],
        font_size=9.1,
    )

    add_heading(doc, "5. Hermes 功能清单", 1)
    add_table(
        doc,
        ["能力域", "主要功能", "关键输出"],
        [
            ["会话与意图", "新建/继续聊天会话、槽位跟踪、短确认解析、覆盖与清空约束", "UserRequirement、questions、dialogue_status"],
            ["集群感知", "读取节点、在线状态、资源余量、路径画像、模型状态", "cluster state、report"],
            ["策略生成", "构造任务画像、过滤不可行节点、生成推荐策略", "policy、task、decision"],
            ["调度评分", "九项指标归一化、多目标权重、硬约束、软偏好", "weights、metric_scores、total_score"],
            ["仿真与解释", "预演推荐策略，给出时延、成本、SLA、安全与风险", "simulation、expected_effect、explanation"],
            ["正式下发", "显式确认后提交策略任务或调度 pending 任务", "lease、submitted_task"],
            ["执行回传", "上报进度、结果、成本、错误、健康度与性能因子", "progress、execution record"],
            ["反馈优化", "解析用户自然语言反馈，生成新的指标偏好并重算策略", "preference_delta、optimized policy"],
            ["外部接入", "Dashboard、HTTP、MCP、CloudSimPlus、真实节点 Agent", "统一工具契约与状态"],
        ],
        [1700, 4200, 3460],
        font_size=9.0,
    )

    add_heading(doc, "6. 代码变更范围", 1)
    add_table(
        doc,
        ["文件", "本次变更"],
        [
            ["src/tianjun/core/policy.py", "为 UserRequirement 增加 priority_vector、metric_preferences、slot_confidence；扩展反馈目标类型。"],
            ["src/tianjun/policy/generator.py", "扩展隐式业务识别、地域别名和模糊纠错；新增向量映射、槽位置信度与反馈累积。"],
            ["src/tianjun/policy/feedback.py", "修复反馈键脱节，增加真实 METRIC_KEYS 映射、反馈维度与幅度感知。"],
            ["src/tianjun/policy/clarifier.py", "追问按调度影响排序；地域不可用时优先询问是否放宽。"],
            ["src/tianjun/scheduling/engine.py", "在 _derive_task_weights() 中叠加 task.intent_weights。"],
            ["src/tianjun/domain/task.py", "任务结构增加 intent_weights，并进入序列化。"],
            ["src/tianjun/scenarios/fixtures.py", "恢复任务时读取 intent_weights，保持持久化兼容。"],
            ["src/tianjun/chat/runtime.py", "LLM 槽位 schema 支持 priority_vector，并做范围校验。"],
            ["src/tianjun/application/control_plane.py", "追问接入在线节点上下文；新增反馈目标进入优化流程。"],
            ["src/tianjun/domain/node.py", "补充城市到 service region 的映射。"],
            ["tests/test_hermes_optimization.py", "新增七组 Hermes 优化回归用例。"],
        ],
        [3000, 6360],
        font_size=8.9,
    )

    add_heading(doc, "7. 验证结果", 1)
    add_callout(
        doc,
        "自动化验证",
        "已安装 pytest 开发依赖并运行 python -m pytest：10 passed in 0.08s。另运行 python -m compileall -q src tests，未发现语法错误。",
        fill="EAF4EE",
        accent=GREEN,
    )
    add_table(
        doc,
        ["验证项", "覆盖内容", "结果"],
        [
            ["隐式工作负载", "在线问答、embedding、批量打标签", "通过"],
            ["地域识别", "苏州、cd、成嘟纠错、武汉", "通过"],
            ["优先级向量", "低时延且不能超预算进入多个真实评分指标", "通过"],
            ["槽位置信度", "缺失推理时延时置信度下降，显式补充后提升", "通过"],
            ["反馈调权", "反馈键均落入真实 METRIC_KEYS，程度词形成不同幅度", "通过"],
            ["反馈落地", "balance 偏好进入 Task.intent_weights", "通过"],
            ["主动追问", "地域无在线节点时优先询问放宽地域", "通过"],
            ["已有拓扑测试", "物理邻居、service region 过滤", "通过"],
        ],
        [1800, 5960, 1600],
        font_size=9.1,
    )

    add_heading(doc, "8. 后续研究建议", 1)
    add_para(
        doc,
        "本次优化完成了可运行、可解释的工程底座。若后续需要形成论文或实验章节，建议把重点放在能够量化比较的方向，而不是继续扩充规则数量。",
    )
    for item in (
        "主动澄清的信息增益实验：对每个待询问槽位做反事实填充，比较候选节点集合、排名稳定性、SLA 达标率和对话轮数。",
        "priority_vector 学习：对比规则映射、小分类模型、LLM few-shot 和人工标注向量，评估调度收益与解释一致性。",
        "slot_confidence 校准：使用历史会话与最终调度结果做可靠性图、Brier Score 或 Expected Calibration Error 对比。",
        "反馈幅度个性化：研究用户对“稍微”“太慢”“无法接受”等表达的个体差异，并使用历史偏好做在线校准。",
        "真实遥测替换：逐步用真实探测时延、链路利用率和节点监控替换当前仿真或合成特征，验证模型分布偏移策略。",
    ):
        add_list_item(doc, item, num_id=research_numbers)

    add_heading(doc, "9. 最终结论", 1)
    add_para(
        doc,
        "Hermes 的最佳优化方向不是把更多决定交给 LLM，而是让自然语言意图以结构化、连续、可校验的方式进入确定性调度内核。本次改动完成了这一关键连接：用户的复合目标能够进入真实评分，反馈能够落到真实指标，置信度能够反映任务差异，追问能够优先处理最影响可行性的约束。与此同时，原有的工具安全边界、仿真预演和人工确认机制保持不变。",
        size=11,
        color=NAVY,
        after=10,
    )
    add_para(doc, "文档结束", size=9, color=MUTED, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.core_properties.title = "Hermes 使用设计评估与优化说明"
    doc.core_properties.subject = "天钧引擎 Hermes 优化评估、实际落地改动与创新功能说明"
    doc.core_properties.author = "Tianjun Optimization Workspace"
    doc.core_properties.keywords = "Hermes, Tianjun, 算力网络, 调度, 意图理解, 反馈闭环"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
